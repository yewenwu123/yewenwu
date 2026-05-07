from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / "澳洲眼镜采购网站执行手册.docx"

ACCENT = RGBColor(184, 106, 52)
DEEP = RGBColor(41, 72, 62)
MUTED = RGBColor(102, 95, 86)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size=11, bold=False, color=None, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
      run.font.color.rgb = color


def add_paragraph(doc, text, style=None, bold=False, color=None, size=11, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=17 if level == 1 else 13, bold=True, color=DEEP, east_asia="Microsoft YaHei UI")
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        if widths:
            hdr_cells[idx].width = widths[idx]
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[idx], "E9D5C3")
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10.5, bold=True, color=DEEP)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if widths:
                cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            set_run_font(run, size=10, color=MUTED)

    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.15)
section.right_margin = Cm(2.15)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

for style_name in ("Heading 1", "Heading 2", "Heading 3"):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")

# Cover
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(80)
title.paragraph_format.space_after = Pt(14)
r = title.add_run("澳洲眼镜采购网站执行手册")
set_run_font(r, size=24, bold=True, color=DEEP, east_asia="Microsoft YaHei UI")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(20)
r = subtitle.add_run("静态展示站 + 咨询表单 + GitHub Pages 免费托管 + 100 美金广告计划")
set_run_font(r, size=13, color=ACCENT)

intro_box = doc.add_table(rows=1, cols=1)
intro_box.autofit = False
intro_box.columns[0].width = Cm(14.8)
cell = intro_box.cell(0, 0)
set_cell_shading(cell, "F7EFE6")
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(
    "这份手册的目标不是让你“先做再说”，而是让你按顺序完成：站点上线、内容替换、表单测试、广告试投、数据复盘。"
)
set_run_font(run, size=11.5, color=MUTED)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(24)
r = meta.add_run("项目文件夹：E:\\ye\\眼镜采购")
set_run_font(r, size=10.5, color=MUTED)

doc.add_page_break()

add_heading(doc, "1. 先明确这次项目要解决什么", level=1)
add_paragraph(doc, "这个网站不是普通零售站，而是一个“建立信任 + 收集咨询”的展示入口。", bold=True, color=DEEP, size=11.5)
add_bullet(doc, "目标客户分成两类：个人买家，以及澳大利亚本地眼镜零售商/卖家。")
add_bullet(doc, "网站不追求当场下单，而是让客户先留下需求，再进入沟通。")
add_bullet(doc, "你的成本重点应放在广告和沟通，不放在网站系统和程序维护。")
add_bullet(doc, "因此本项目采用纯静态网站，适合 GitHub Pages 免费托管。")

add_heading(doc, "2. 当前网站已经包含哪些内容", level=1)
add_table(
    doc,
    ["模块", "作用", "你后面主要改什么"],
    [
        ["首页首屏", "先说明你做的是中国到澳洲的眼镜采购服务", "改标题、改一句话介绍、改按钮文案"],
        ["个人介绍", "让客户知道背后是真人，先建立信任", "放你的生活照、个人故事、做事方式"],
        ["款式展示", "把淘宝或其他渠道找到的款式做成“可咨询”的样式卡片", "换图片、改标题、改描述"],
        ["咨询表单", "收集客户名字、邮箱、WhatsApp、需求", "接入表单服务、测试能否收到表单"],
        ["常见问题", "提前解释为什么不是直接零售下单", "根据你的实际业务补充回答"],
    ],
    widths=[Cm(3.2), Cm(6.1), Cm(6.0)],
)

add_heading(doc, "3. 你最常改的文件在哪", level=1)
add_table(
    doc,
    ["文件路径", "用途", "修改频率"],
    [
        ["data/site-config.js", "联系人姓名、邮箱、WhatsApp、主文案、表单地址", "很高"],
        ["data/products.js", "样式卡片的数据列表", "很高"],
        ["assets/images/", "生活照和款式图", "很高"],
        ["index.html", "页面结构", "中等"],
        ["styles.css", "颜色、排版、按钮、卡片样式", "中等"],
    ],
    widths=[Cm(4.8), Cm(7.0), Cm(3.2)],
)

add_heading(doc, "4. 你以后怎么改生活照片", level=1)
add_bullet(doc, "打开文件夹：assets/images/")
add_bullet(doc, "把你的真实照片准备好，建议横竖比例尽量接近当前占位图。")
add_bullet(doc, "最省事的方式：保留原文件名，直接替换掉 life-photo-1.svg、life-photo-2.svg、life-photo-3.svg 对应的图片。")
add_bullet(doc, "如果你想用 jpg 或 png，也可以把文件放进去，然后到 index.html 里把图片路径改掉。")
add_bullet(doc, "优先放这三种照片：日常生活照、工作或出行照、看起来亲和的个人半身照。")
add_paragraph(doc, "不要只放产品图而不放真人图。跨境采购最怕不信任，真人感比“精修品牌感”更重要。", color=ACCENT, bold=True, size=11)

add_heading(doc, "5. 你以后怎么上传淘宝款式", level=1)
add_paragraph(doc, "网站的款式区不是后台系统，而是一个简单的数据文件，所以你自己改也不会乱。", color=MUTED, size=10.8)
add_bullet(doc, "把款式图片放进 assets/images/，比如命名为 style-4.jpg、style-5.jpg。")
add_bullet(doc, "打开 data/products.js。每一组大括号就是一个产品卡片。")
add_bullet(doc, "复制一组现有内容，改 5 个地方：title、category、audience、image、description。")
add_paragraph(doc, "示例字段解释：", bold=True, color=DEEP)
add_table(
    doc,
    ["字段", "示例", "说明"],
    [
        ["title", "Soft Rectangle Acetate", "客户看到的款式名称"],
        ["category", "Acetate", "材质或风格标签"],
        ["audience", "Boutique retail", "更适合哪类客户"],
        ["image", "assets/images/style-4.jpg", "图片路径"],
        ["description", "A clean everyday frame...", "简短描述，不要写得像淘宝详情页"],
    ],
    widths=[Cm(2.7), Cm(5.8), Cm(6.5)],
)

add_heading(doc, "6. 表单怎么接起来，尽量不花钱", level=1)
add_paragraph(doc, "推荐顺序：先用免费表单服务把咨询收起来，等网站稳定了再考虑更复杂的 CRM。", color=MUTED)
add_table(
    doc,
    ["方案", "适合你吗", "优点", "注意点"],
    [
        ["FormSubmit", "很适合起步", "不用自己写后端，适合静态网站，接入快", "第一次要验证邮箱，功能偏简单"],
        ["Formspree", "也适合", "界面更完整，适合后面稍微规范化管理", "免费额度有限，更适合低咨询量起步"],
    ],
    widths=[Cm(3.0), Cm(3.5), Cm(5.1), Cm(5.0)],
)
add_bullet(doc, "你要改的地方只有一个：data/site-config.js 里的 formEndpoint。")
add_bullet(doc, "例如：如果表单服务给你一个提交地址，就把那个地址填进去。")
add_bullet(doc, "填完之后一定要自己试填 2 次：一次正常内容，一次只填最少字段，确认邮箱能收到。")

add_heading(doc, "7. 怎么免费托管到 GitHub Pages", level=1)
add_bullet(doc, "新建一个 GitHub 仓库，把整个项目文件夹上传。")
add_bullet(doc, "仓库里要保持 index.html 在根目录。")
add_bullet(doc, "进入 GitHub 仓库的 Settings -> Pages。")
add_bullet(doc, "Source 选择：Deploy from a branch。")
add_bullet(doc, "Branch 选择 main，Folder 选择 /root。")
add_bullet(doc, "保存后等几分钟，GitHub 会生成一个公开网址。")
add_bullet(doc, "如果后面你买域名，再把自定义域名绑上去；但起步阶段不强制。")
add_paragraph(doc, "建议先用 GitHub Pages 的免费二级域名把流程跑通，不要一上来就花钱买太多工具。", color=ACCENT, bold=True)

add_heading(doc, "8. 100 美金广告预算，不要盲目投", level=1)
add_paragraph(doc, "如果你每个月只有 100 美金广告预算，最重要的不是“同时投很多平台”，而是先让一个渠道跑顺。", bold=True, color=DEEP, size=11.3)
add_table(
    doc,
    ["阶段", "预算", "动作", "目标"],
    [
        ["第 0 周", "$0", "先不上广告，先补真人照片、补 6-9 个款式、表单测试通过", "确保网站能接咨询"],
        ["第 1-4 周", "$90", "只投一个主渠道，建议先投 Meta/Instagram 图文广告", "看哪个图片和文案最能让人咨询"],
        ["机动预算", "$10", "给表现最好的广告素材补一点预算", "做小范围验证，不乱加钱"],
    ],
    widths=[Cm(2.6), Cm(2.2), Cm(7.0), Cm(4.2)],
)

add_heading(doc, "9. 为什么建议先投 Meta，而不是一开始就投很多平台", level=1)
add_bullet(doc, "眼镜是视觉类产品，图片和生活感更容易在 Meta / Instagram 里建立初步兴趣。")
add_bullet(doc, "你现在网站不是电商结账页，而是咨询页，更适合“引导留言或表单提交”的广告目标。")
add_bullet(doc, "预算太小的时候，同时投 Meta 和 Google Ads，往往每个平台都不够学习。")
add_bullet(doc, "如果以后你发现很多人会主动搜“wholesale eyewear australia”这类词，再增加 Google Search。")

add_heading(doc, "10. 第一个月广告怎么做", level=1)
add_table(
    doc,
    ["项目", "建议做法"],
    [
        ["广告目标", "优先选引导咨询、消息或网站提交表单，不要一开始追求购买转化"],
        ["投放地区", "先选澳大利亚主要城市：Sydney、Melbourne、Brisbane、Perth、Adelaide"],
        ["素材数量", "至少准备 3 套素材：1 套真人感、1 套款式拼图、1 套“代采购流程说明”"],
        ["文案重点", "不要只写便宜，要写“可先沟通、可看款式、可服务个人和商家、不是盲买”"],
        ["落地页按钮", "统一导向咨询表单，不要做复杂分流"],
    ],
    widths=[Cm(3.0), Cm(12.5)],
)

add_heading(doc, "11. 广告文案方向建议", level=1)
add_bullet(doc, "个人客户方向：Looking for unique frames sourced from China? Start with a conversation, not a rushed checkout.")
add_bullet(doc, "商家客户方向：Need a sourcing contact for eyewear styles from China? I help Australian sellers review styles before ordering.")
add_bullet(doc, "共同重点：先咨询、再确认、再采购。")

add_heading(doc, "12. 每周执行节奏，避免乱做", level=1)
add_table(
    doc,
    ["时间", "你要做什么", "完成标准"],
    [
        ["周一", "新增 2-3 个款式，检查网站文案和图片", "站内内容看起来持续在更新"],
        ["周二", "自己提交一次测试表单", "确认邮箱/表单服务正常"],
        ["周三", "检查广告点击和咨询量", "记录哪套图点击高，哪套图有人问"],
        ["周四", "优化 1 条主文案或 1 组素材", "每周只改一个重点，不要全改"],
        ["周五", "整理客户咨询内容", "记录客户最常问的问题，下周补到 FAQ"],
    ],
    widths=[Cm(2.3), Cm(8.0), Cm(5.2)],
)

add_heading(doc, "13. 不要做的事", level=1)
add_bullet(doc, "不要一开始就做成几百个 SKU 的零售商城。")
add_bullet(doc, "不要把有限预算分给太多广告平台。")
add_bullet(doc, "不要用太多花哨功能，结果表单没接通。")
add_bullet(doc, "不要长期只放产品图，不放你的真人信息。")
add_bullet(doc, "不要改一堆变量后却不做测试。每次改完至少自己点一遍。")

add_heading(doc, "14. 你现在最应该先做的 10 件事", level=1)
for item in [
    "把你的真实姓名、邮箱、WhatsApp 改进 data/site-config.js。",
    "准备 3 张生活照，替换 assets/images/ 里的占位图。",
    "先准备 6 到 9 个你最有把握的款式，不要贪多。",
    "把这些款式图放进 assets/images/ 并更新 data/products.js。",
    "选一个免费表单服务，把 formEndpoint 填进去。",
    "自己做 2 次表单测试，确认你能收到咨询。",
    "把项目上传到 GitHub 仓库。",
    "打开 GitHub Pages，让网站有公开链接。",
    "给 3 套广告素材各写一条英文文案。",
    "网站上线后再开始投 100 美金预算，不要倒过来。"
]:
    add_bullet(doc, item)

add_heading(doc, "15. 附录：本项目里已经为你准备好的内容", level=1)
add_bullet(doc, "适合澳洲客户浏览的英文页面结构。")
add_bullet(doc, "生活照占位图，提醒你该放什么类型的真人图片。")
add_bullet(doc, "样式卡片数据文件，方便你自己追加淘宝款式。")
add_bullet(doc, "适合 GitHub Pages 的纯静态结构。")
add_bullet(doc, "一个低成本咨询表单入口。")

add_heading(doc, "16. 参考链接", level=1)
add_bullet(doc, "GitHub Pages 官方文档：https://docs.github.com/en/pages")
add_bullet(doc, "GitHub Pages 快速开始：https://docs.github.com/en/pages/getting-started-with-github-pages")
add_bullet(doc, "Formspree 官网定价页：https://formspree.io/plans")
add_bullet(doc, "FormSubmit 官网：https://formsubmit.co/")
add_bullet(doc, "Google Ads Keyword Planner 官方帮助：https://support.google.com/google-ads/answer/7337243")

doc.save(OUT_PATH)
print(OUT_PATH)
